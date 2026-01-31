# src/report_sections/conclusion.py
from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from src.report_sections._word_common import (
    s,
    body,
    bullets_from_text,
    add_heading,                    # TOC-safe heading
    set_paragraph_bottom_border,    # orange underline
    tight_paragraph,
)

# -------------------------
# Section style (match your new rule)
# -------------------------
TITLE_FONT = "Cambria"
TITLE_SIZE = 16  # ✅ Heading 1 titles = 16
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"


def _heading_with_orange_line(
    doc: Document,
    text: str,
    *,
    level: int,
    font: str = TITLE_FONT,
    size: int = TITLE_SIZE,
    color: Optional[RGBColor] = TITLE_BLUE,
    orange_hex: str = ORANGE_HEX,
    after_pt: float = 6,
) -> None:
    """
    Create a REAL Word Heading (TOC-safe) and add an orange bottom border under it.
    """
    p = add_heading(
        doc,
        s(text),
        level=int(level),
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font=font,
        size=int(size),
        bold=True,
        color=color,
    )

    # consistent spacing
    tight_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before_pt=0,
        after_pt=float(after_pt),
        line_spacing=1,
    )

    # orange underline (bottom border)
    set_paragraph_bottom_border(
        p,
        color_hex=orange_hex,
        size_eighths=12,  # 12 => 1.5pt
        space=2,
    )


def add_conclusion_section(
    doc: Document,
    *,
    conclusion_text: Optional[str] = None,
    key_points: Optional[List[str]] = None,
    recommendations_summary: Optional[str] = None,
    section_no: str = "7",
) -> None:
    """
    Final section: Conclusion (TOC-safe)
    - Heading 1 for main title (shows in TOC)
    - Heading 2 for subsections (shows in TOC when TOC levels include 1-3)
    """
    doc.add_page_break()

    # ✅ Heading 1 + size 16 + orange underline
    _heading_with_orange_line(
        doc,
        f"{s(section_no)}.        Conclusion:",
        level=1,
        after_pt=6,
    )

    # 1) Main conclusion paragraph (always print something)
    main_text = s(conclusion_text)
    if main_text:
        body(doc, main_text)
    else:
        body(
            doc,
            "Overall, the monitoring confirmed that the assessed WASH intervention is functional and "
            "providing services to the beneficiary community. Addressing the observed technical and "
            "operational gaps through timely corrective actions and strengthened O&M capacity will "
            "improve system reliability and long-term sustainability."
        )

    # small spacing after main paragraph (controlled)
    doc.add_paragraph("")

    # 2) Key points (optional)
    kp = [s(x) for x in (key_points or []) if s(x)]
    if kp:
        add_heading(
            doc,
            "Key Points",
            level=2,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font="Cambria",
            size=12,
            bold=True,
            color=None,  # black
        )
        tight_paragraph(doc.paragraphs[-1], before_pt=0, after_pt=4, line_spacing=1)

        for it in kp:
            body(doc, f"• {it}")

        doc.add_paragraph("")

    # 3) Recommendations summary (optional)
    rec_text = s(recommendations_summary)
    if rec_text:
        add_heading(
            doc,
            "Recommendations Summary",
            level=2,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            font="Cambria",
            size=12,
            bold=True,
            color=None,  # black
        )
        tight_paragraph(doc.paragraphs[-1], before_pt=0, after_pt=4, line_spacing=1)

        items = bullets_from_text(rec_text)
        if items:
            for it in items:
                body(doc, f"• {it}")
        else:
            body(doc, rec_text)

        doc.add_paragraph("")


# ✅ Compatibility wrapper (some builders may call add_conclusion)
def add_conclusion(
    doc: Document,
    *,
    row: Optional[Dict[str, Any]] = None,
    conclusion_text: Optional[str] = None,
    key_points: Optional[List[str]] = None,
    recommendations_summary: Optional[str] = None,
    component_observations: Optional[List[Dict[str, Any]]] = None,
    severity_by_no: Any = None,
    section_no: str = "7",
    **kwargs,
) -> None:
    """
    Compatibility wrapper so existing pipelines keep working.
    Extra params are accepted but not required here.
    """
    _ = row, component_observations, severity_by_no, kwargs

    add_conclusion_section(
        doc,
        conclusion_text=conclusion_text,
        key_points=key_points,
        recommendations_summary=recommendations_summary,
        section_no=section_no,
    )
