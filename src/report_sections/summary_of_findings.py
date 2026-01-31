# src/report_sections/summary_of_findings.py
from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx.document import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor

from src.report_sections._word_common import (
    s,
    # text/paragraph
    tight_paragraph,
    set_run,

    # ✅ Standard H1 title (size 16 + TOC + orange line)
    add_section_title_h1,

    # table sizing/layout
    set_table_fixed_layout,
    set_table_width_from_section,
    set_table_borders,
    shade_cell,
    set_cell_borders,
    set_cell_margins,
    set_row_cant_split,
    set_row_height_exact,
    set_row_height_at_least,
    section_usable_width_emu,

    # section-6 shared helpers
    normalize_sentence,
    severity_checkbox_line,
    extract_findings_and_recos_from_section5,
    present_severities_from_mapping,
)

# -------------------------
# Style (keep section-specific only)
# -------------------------
TITLE_TEXT = "6.        Summary of the findings:"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"

DARK_BLUE_HEX = "1F4E79"
BORDER_HEX = "000000"
HEADER_TEXT_COLOR = RGBColor(255, 255, 255)

FONT_BODY = "Times New Roman"
FONT_SIZE_BODY = 10
FONT_SIZE_HEADER = 10


def add_summary_of_findings_section6(
    doc: Document,
    *,
    component_observations: List[Dict[str, Any]],
    severity_by_no: Optional[Dict[int, str]] = None,
    severity_by_finding: Optional[Dict[str, str]] = None,
    add_legend: bool = True,
) -> None:
    """
    Section 6:
      - Summary of the findings table
      - Legend table (optional)
    """
    severity_by_no = severity_by_no or {}
    severity_by_finding = severity_by_finding or {}

    doc.add_page_break()

    # ✅ Proper Heading 1 + orange underline on SAME paragraph (no extra blank title line)
    add_section_title_h1(
        doc,
        TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=TITLE_BLUE,
        orange_hex=ORANGE_HEX,
        after_pt=6,
    )
    doc.add_paragraph("")

    extracted = extract_findings_and_recos_from_section5(component_observations or [])
    if not extracted:
        p = doc.add_paragraph()
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        r = p.add_run("No major findings were captured in Section 5 to summarize.")
        set_run(r, "Times New Roman", 11, bold=False)
        return

    # usable width (for “remaining column width” calculation)
    usable_emu = section_usable_width_emu(doc.sections[0])
    usable_in = usable_emu / 914400.0  # EMU -> inch

    # ---- Main table
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    set_table_fixed_layout(table)
    set_table_width_from_section(table, doc, section_index=0)
    set_table_borders(table, color_hex=BORDER_HEX, size="8")

    # Column widths (inches; stable)
    w_no = 0.45
    w_find = 4.20
    w_sev = 0.80
    w_rec = max(0.90, usable_in - (w_no + w_find + w_sev))

    widths = [Inches(w_no), Inches(w_find), Inches(w_sev), Inches(w_rec)]
    headers = ["No.", "Finding", "Severity", "Recommendation\n/ Corrective\nAction"]

    # Header row
    header_row = table.rows[0]
    set_row_cant_split(header_row, cant_split=True)
    set_row_height_exact(header_row, 0.22)

    for i, cell in enumerate(header_row.cells):
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""

        shade_cell(cell, DARK_BLUE_HEX)
        set_cell_borders(cell, size=10, color_hex=BORDER_HEX)
        set_cell_margins(cell, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        p = cell.paragraphs[0]
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)
        rr = p.add_run(headers[i])
        set_run(rr, FONT_BODY, FONT_SIZE_HEADER, True, HEADER_TEXT_COLOR)

    # Body rows
    for idx, item in enumerate(extracted, start=1):
        finding = normalize_sentence(item.get("finding", ""))
        reco = normalize_sentence(item.get("recommendation", "")) or "—"

        row = table.add_row()

        # ✅ IMPORTANT: allow split to avoid huge whitespace when finding text is long
        set_row_cant_split(row, cant_split=False)

        set_row_height_at_least(row, 0.36)

        cells = row.cells
        for i, c in enumerate(cells):
            c.width = widths[i]
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            c.text = ""
            set_cell_borders(c, size=8, color_hex=BORDER_HEX)
            set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        # No
        p0 = cells[0].paragraphs[0]
        tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p0.add_run(str(idx)), FONT_BODY, FONT_SIZE_BODY, False)

        # Finding
        p1 = cells[1].paragraphs[0]
        tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p1.add_run(finding), FONT_BODY, FONT_SIZE_BODY, False)

        # Severity (checkbox line)
        chosen = ""
        if idx in severity_by_no:
            chosen = s(severity_by_no.get(idx))
        else:
            # match by normalized finding
            f_norm = finding.lower()
            for k, v in severity_by_finding.items():
                if normalize_sentence(k).lower() == f_norm:
                    chosen = s(v)
                    break

        p2 = cells[2].paragraphs[0]
        tight_paragraph(p2, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p2.add_run(severity_checkbox_line(chosen)), FONT_BODY, 9, False)

        # Recommendation
        p3 = cells[3].paragraphs[0]
        tight_paragraph(p3, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p3.add_run(reco), FONT_BODY, FONT_SIZE_BODY, False)

    doc.add_paragraph("")

    # ---- Legend
    if not add_legend:
        return

    p_leg = doc.add_paragraph()
    tight_paragraph(p_leg, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
    set_run(p_leg.add_run("Legend:"), "Times New Roman", 11, True)
    doc.add_paragraph("")

    legend_definitions = {
        "High": "Critical issue affecting functionality, safety, or compliance; requires immediate action.",
        "Medium": "Moderate issue affecting efficiency or performance; corrective action required.",
        "Low": "Minor issue with limited impact; corrective action recommended.",
    }

    present = present_severities_from_mapping(extracted, severity_by_no, severity_by_finding)
    if not present:
        present = ["High", "Medium", "Low"]

    legend = doc.add_table(rows=1, cols=2)
    legend.autofit = False
    legend.alignment = WD_TABLE_ALIGNMENT.LEFT
    legend.style = "Table Grid"

    set_table_fixed_layout(legend)
    set_table_width_from_section(legend, doc, section_index=0)
    set_table_borders(legend, color_hex=BORDER_HEX, size="8")

    # Legend header
    r0 = legend.rows[0]
    set_row_cant_split(r0, cant_split=True)
    set_row_height_exact(r0, 0.22)

    h0, h1c = r0.cells
    for c, txt in ((h0, "Severity"), (h1c, "Definition")):
        c.text = ""
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c, DARK_BLUE_HEX)
        set_cell_borders(c, size=10, color_hex=BORDER_HEX)
        set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        pp = c.paragraphs[0]
        tight_paragraph(pp, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)
        set_run(pp.add_run(txt), FONT_BODY, 10, True, HEADER_TEXT_COLOR)

    # Legend rows
    for sev in present:
        rr = legend.add_row()

        # ✅ allow split (legend definitions can wrap)
        set_row_cant_split(rr, cant_split=False)

        set_row_height_at_least(rr, 0.25)

        c0, c1 = rr.cells
        for c in (c0, c1):
            c.text = ""
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_borders(c, size=8, color_hex=BORDER_HEX)
            set_cell_margins(c, top_dxa=80, bottom_dxa=80, left_dxa=120, right_dxa=120)

        p0 = c0.paragraphs[0]
        tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p0.add_run(sev), FONT_BODY, 10, False)

        p1 = c1.paragraphs[0]
        tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p1.add_run(legend_definitions.get(sev, "")), FONT_BODY, 10, False)
