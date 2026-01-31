# src/report_sections/data_collection_methods.py
from __future__ import annotations

from typing import Any, Dict, Optional, List

from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from src.report_sections._word_common import (
    s,
    tight_paragraph,
    set_run,
    parse_bool_like,

    # ✅ standard H1 title (size 16 + TOC + orange line)
    add_section_title_h1,
)

# -------------------------
# Section style
# -------------------------
TITLE_TEXT = "3.        Data Collection Methods:"
TITLE_FONT = "Cambria"
TITLE_SIZE = 16                 # ✅ as requested
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"

BODY_FONT = "Times New Roman"
BODY_SIZE = 11

# ✅ Two-line gap after this section
AFTER_SECTION_GAP_PT = 24


# -------------------------
# Internal helpers
# -------------------------
def _yes(v: Any) -> bool:
    return parse_bool_like(v) is True


def _pick(row: Dict[str, Any], overrides: Dict[str, Any], *keys: str) -> Any:
    """
    Pick first non-empty value from overrides then row.
    (kept local for speed; avoids importing extra helper if not needed)
    """
    for k in keys:
        if k in overrides and overrides.get(k) not in (None, "", " "):
            return overrides.get(k)
    for k in keys:
        if row.get(k) not in (None, "", " "):
            return row.get(k)
    return None


def _add_numbered_item(doc: Document, text: str) -> None:
    """Add a clean Word numbered list item."""
    p = doc.add_paragraph(style="List Number")
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1.0)
    set_run(p.add_run(s(text)), BODY_FONT, BODY_SIZE, bold=False)


def _add_two_line_gap(doc: Document) -> None:
    """
    Robust 2-line visual gap before next section starts (if next section doesn't force a page break).
    Uses spacing-after instead of multiple empty paragraphs to avoid collapsing.
    """
    p = doc.add_paragraph("")
    tight_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before_pt=0,
        after_pt=float(AFTER_SECTION_GAP_PT),
        line_spacing=1.0,
    )


# -------------------------
# Main section
# -------------------------
def add_data_collection_methods(
    doc: Document,
    row: Dict[str, Any],
    overrides: Optional[Dict[str, Any]] = None,
) -> None:
    row = row or {}
    overrides = overrides or {}

    doc.add_page_break()

    # ✅ Standard Section Title: Heading 1 + size 16 + orange underline + TOC safe
    add_section_title_h1(
        doc,
        TITLE_TEXT,
        font=TITLE_FONT,
        size=TITLE_SIZE,
        color=TITLE_BLUE,
        orange_hex=ORANGE_HEX,
        after_pt=6,
    )

    # =========================================================
    # 1. Detect available evidence (prefer overrides, fallback row)
    # =========================================================
    has_contract = _yes(_pick(row, overrides, "D1_contract_available"))
    has_journal = _yes(_pick(row, overrides, "D1_journal_available"))
    has_boq = _yes(_pick(row, overrides, "D2_boq_available"))
    has_drawings = _yes(_pick(row, overrides, "D2_drawings_available"))
    has_geo_tests = _yes(_pick(row, overrides, "D3_geophysical_tests_available"))
    has_wq_tests = _yes(_pick(row, overrides, "D4_water_quality_tests_available"))
    has_pump_tests = _yes(_pick(row, overrides, "D4_pump_test_results_available"))

    # Data-collection actions
    has_observation = _yes(_pick(row, overrides, "D0_direct_observation"))
    has_interview = _yes(_pick(row, overrides, "D0_key_informant_interview"))
    has_photos = _yes(_pick(row, overrides, "D0_photos_taken"))
    has_gps = _yes(_pick(row, overrides, "D0_gps_points_recorded"))

    # =========================================================
    # 2. Build DOCUMENT REVIEW text dynamically
    # =========================================================
    reviewed_docs: List[str] = []

    if has_boq:
        reviewed_docs.append("Bill of Quantities (BOQ)")
    if has_drawings:
        reviewed_docs.append("approved technical drawings")
    if has_contract:
        reviewed_docs.append("contract documents")
    if has_journal:
        reviewed_docs.append("site journal and progress records")
    if has_geo_tests:
        reviewed_docs.append("geophysical and hydrological test reports")
    if has_wq_tests:
        reviewed_docs.append("water quality test results")
    if has_pump_tests:
        reviewed_docs.append("pump test results")

    doc_review_phrase = ""
    if reviewed_docs:
        doc_review_phrase = "Review of project documentation, including " + ", ".join(reviewed_docs) + "."

    # =========================================================
    # 3. Build numbered METHODS list (automatic)
    # =========================================================
    methods: List[str] = []

    if has_observation:
        methods.append("Direct technical observation of work progress and construction quality on-site.")

    if doc_review_phrase:
        methods.append(doc_review_phrase)

    if has_interview:
        methods.append(
            "Semi-structured interviews with technical staff of the contracted company, implementing partner personnel, "
            "and Community Development Council (CDC) members."
        )

    if has_photos:
        methods.append(
            "Collection and review of geo-referenced photographic evidence to verify physical progress and workmanship."
        )

    if has_gps:
        methods.append(
            "Verification of GPS coordinates and location data to confirm site positioning and component alignment."
        )

    # Fallback safety
    if not methods:
        methods.append(
            "The monitoring visit applied standard Third-Party Monitoring (TPM) data collection techniques in line with UNICEF WASH guidelines."
        )

    for m in methods:
        _add_numbered_item(doc, m)

    # small spacing before narrative
    p_gap = doc.add_paragraph("")
    tight_paragraph(p_gap, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=6, line_spacing=1)

    # =========================================================
    # 4. Professional narrative paragraph (automatic)
    # =========================================================
    narrative = (
        "The Third-Party Monitoring (TPM) assessment was conducted using a structured mixed-methods approach, combining "
        "direct on-site technical observation, systematic review of available project documentation, and qualitative "
        "engagement with relevant stakeholders. The monitoring focused on verifying construction quality, system "
        "functionality, and compliance with approved designs and contractual requirements, while identifying technical "
        "and operational risks that may affect performance and sustainability. Physical and documentary evidence was "
        "assessed across all applicable project components, and findings were analyzed, categorized by severity, and "
        "linked to practical corrective actions in accordance with UNICEF WASH standards and third-party monitoring protocols."
    )

    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before_pt=0, after_pt=0, line_spacing=1.15)
    set_run(p.add_run(narrative), BODY_FONT, BODY_SIZE, bold=False)

    # ✅ Guaranteed 2-line gap before next content (if next section doesn't force a new page)
    _add_two_line_gap(doc)
