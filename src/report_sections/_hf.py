# src/report_sections/_hf.py
from __future__ import annotations

from PIL import Image

from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _clear_paragraph(p) -> None:
    p.text = ""
    for r in p.runs:
        r.text = ""


def _tight_paragraph(p, align) -> None:
    p.alignment = align
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    p.paragraph_format.line_spacing = 1


def _set_table_fixed_layout(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _remove_all_tables_from_footer(footer) -> None:
    for tbl in list(footer._element.xpath(".//w:tbl")):
        tbl.getparent().remove(tbl)


def _img_aspect_ratio(path: str) -> float:
    """returns width/height"""
    with Image.open(path) as im:
        w, h = im.size
    return (w / h) if h else 1.0


def _safe_add_picture(paragraph, path, *, width=None, height=None) -> None:
    if not path:
        return
    try:
        paragraph.add_run().add_picture(path, width=width, height=height)
    except Exception:
        pass


def apply_header_footer(
    doc,
    *,
    unicef_logo_path: str | None = None,
    act_logo_path: str | None = None,
    ppc_logo_path: str | None = None,

    # Header size
    unicef_width=Inches(1.6),

    # Footer sizes (height-based for easy alignment)
    act_height=Inches(1.40),
    ppc_height=Inches(0.85),

    # Footer layout tuning
    footer_left_pct: float = 0.45,   # ACT column
    footer_mid_pct: float = 0.10,    # spacer column
    ppc_max_width=Inches(1.60),      # PPC never exceeds this width
) -> None:
    """
    Header: UNICEF centered (width-controlled)
    Footer: ACT | spacer | PPC (height-controlled, PPC width capped)
    Applies to ALL sections/pages.

    Optimization:
      - PPC aspect ratio is computed ONCE (Image.open only once),
        instead of once per section.
    """

    # clamp percentages (avoid weird values)
    footer_left_pct = max(0.10, min(float(footer_left_pct), 0.80))
    footer_mid_pct = max(0.00, min(float(footer_mid_pct), 0.40))
    if footer_left_pct + footer_mid_pct > 0.95:
        footer_mid_pct = 0.95 - footer_left_pct

    # ---------------------------
    # PPC sizing decision cached
    # ---------------------------
    # We decide ONCE whether PPC should use width cap or height.
    ppc_ratio: float = 1.0
    ppc_ratio_ok = False

    if ppc_logo_path:
        try:
            ppc_ratio = _img_aspect_ratio(ppc_logo_path)  # w/h
            ppc_ratio_ok = True
        except Exception:
            ppc_ratio_ok = False

    # Pre-compute the best PPC add_picture mode (once)
    # We keep identical behavior:
    #   - estimate width based on height * ratio
    #   - if estimated width exceeds ppc_max_width -> use width cap
    #   - else use height
    ppc_use_width_cap = True  # default fallback safe behavior
    if ppc_logo_path and ppc_ratio_ok:
        try:
            est_width_emu = int(int(ppc_height) * float(ppc_ratio))
            ppc_use_width_cap = est_width_emu > int(ppc_max_width)
        except Exception:
            ppc_use_width_cap = True

    for sec in doc.sections:
        sec.different_first_page_header_footer = False

        # ----------------
        # HEADER (UNICEF)
        # ----------------
        hdr = sec.header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        _clear_paragraph(p)
        _tight_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
        _safe_add_picture(p, unicef_logo_path, width=unicef_width)

        # ----------------
        # FOOTER (ACT | spacer | PPC)
        # ----------------
        ftr = sec.footer
        if ftr.paragraphs:
            _clear_paragraph(ftr.paragraphs[0])
        _remove_all_tables_from_footer(ftr)

        usable_width = sec.page_width - sec.left_margin - sec.right_margin
        usable_width_emu = int(usable_width)

        tbl = ftr.add_table(rows=1, cols=3, width=usable_width)
        tbl.autofit = False
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_table_fixed_layout(tbl)

        w0 = int(usable_width_emu * footer_left_pct)
        w1 = int(usable_width_emu * footer_mid_pct)
        w2 = max(int(usable_width_emu - (w0 + w1)), 1)

        tbl.columns[0].width = w0
        tbl.columns[1].width = w1
        tbl.columns[2].width = w2

        # ACT (left)
        c0 = tbl.cell(0, 0)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        _clear_paragraph(p0)
        _tight_paragraph(p0, WD_ALIGN_PARAGRAPH.LEFT)
        _safe_add_picture(p0, act_logo_path, height=act_height)

        # spacer (middle)
        c1 = tbl.cell(0, 1)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        _clear_paragraph(p1)
        _tight_paragraph(p1, WD_ALIGN_PARAGRAPH.CENTER)

        # PPC (right) with width cap (decision cached)
        c2 = tbl.cell(0, 2)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = c2.paragraphs[0]
        _clear_paragraph(p2)
        _tight_paragraph(p2, WD_ALIGN_PARAGRAPH.RIGHT)

        if ppc_logo_path:
            try:
                if ppc_use_width_cap:
                    _safe_add_picture(p2, ppc_logo_path, width=ppc_max_width)
                else:
                    _safe_add_picture(p2, ppc_logo_path, height=ppc_height)
            except Exception:
                _safe_add_picture(p2, ppc_logo_path, width=ppc_max_width)
