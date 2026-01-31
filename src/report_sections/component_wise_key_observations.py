# src/report_sections/component_wise_key_observations.py
from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx.document import Document
from docx.shared import RGBColor

from src.report_sections._word_common import (
    s,
    # headings/body
    add_section_title_h1,            # ✅ H1 (size 16 + TOC + orange line)
    add_heading,                     # ✅ generic TOC-safe heading for H2/H3
    body,
    strip_heading_numbering,
    special_title_normalized,
    # section 5 blocks
    add_major_findings_table_tool6,
    add_text_left_photo_right_block,
)

# -------------------------------------------------
# Styling
# -------------------------------------------------
BLACK = RGBColor(0, 0, 0)
TITLE_BLUE = RGBColor(0, 112, 192)
ORANGE_HEX = "ED7D31"


def _h2(doc: Document, text: str) -> None:
    """Heading 2 (TOC-safe), default style color."""
    add_heading(
        doc,
        s(text),
        level=2,
        align=None,
        font="Cambria",
        size=13,
        bold=True,
        color=None,  # keep Heading 2 style default (or doc default)
    )


def _h3_black(doc: Document, text: str) -> None:
    """
    Heading 3 and FORCE black color (overrides any style color).
    This matches your requirement: all sub-subtitles are Heading 3 BLACK.
    """
    p = doc.add_paragraph(s(text), style="Heading 3")
    for r in p.runs:
        try:
            r.font.color.rgb = BLACK
        except Exception:
            pass


def add_component_wise_key_observations_tool6(
    doc: Document,
    component_observations: List[Dict[str, Any]],
    photo_bytes: Dict[str, bytes],
    photo_field_map: Optional[Dict[str, str]] = None,
) -> None:
    """
    Section 5: Project Component-Wise Key Observations

    Rules enforced:
      - Section title: Heading 1 (size 16 + orange underline)
      - Component titles: Heading 2
      - Subsection titles: Heading 3 (FORCE black)
    """
    component_observations = component_observations or []
    photo_bytes = photo_bytes or {}
    photo_field_map = photo_field_map or {}

    if not component_observations:
        return

    doc.add_page_break()

    # ✅ Section 5 main title => Heading 1 + orange line (TOC friendly)
    add_section_title_h1(
        doc,
        "5.        Project Component-Wise Key Observations:",
        font="Cambria",
        size=16,
        color=TITLE_BLUE,
        orange_hex=ORANGE_HEX,
        after_pt=6,
    )
    doc.add_paragraph("")

    for comp in component_observations:
        if not isinstance(comp, dict):
            continue

        raw_title = s(comp.get("title"))
        title_clean = raw_title if raw_title else s(comp.get("comp_id"))

        # ---------------- Component heading (H2)
        if title_clean:
            _h2(doc, title_clean)
            doc.add_paragraph("")

        # ---------------- Intro paragraphs
        for para in (comp.get("paragraphs") or []):
            if s(para):
                body(doc, para)
        doc.add_paragraph("")

        comp_id = s(comp.get("comp_id")).strip()
        subs = comp.get("subsections") or []

        # ---------------- Subsections (H3) => ALWAYS BLACK
        for sub_idx, sub in enumerate(subs, start=1):
            if not isinstance(sub, dict):
                continue

            stitle = s(sub.get("title")).strip()
            if stitle:
                # normalize + clean numbering
                _ = special_title_normalized(stitle)  # keep normalization call (future rules)
                clean_title = strip_heading_numbering(stitle).rstrip(":").strip()

                if comp_id:
                    heading_text = f"{comp_id}.{sub_idx} {clean_title}"
                else:
                    heading_text = f"{sub_idx}. {clean_title}"

                # ✅ always H3 black
                _h3_black(doc, heading_text)
                doc.add_paragraph("")

            # ---------- Content
            mt = sub.get("major_table")
            if isinstance(mt, list):
                add_major_findings_table_tool6(
                    doc,
                    major_rows=mt,
                    photo_bytes=photo_bytes,
                    photo_field_map=photo_field_map,
                )
                doc.add_paragraph("")
            else:
                for para in (sub.get("paragraphs") or []):
                    if s(para):
                        body(doc, para)
                doc.add_paragraph("")

        # ---------------- Photos block
        photos = comp.get("photos") or []
        if photos:
            left_text = "\n".join([s(x) for x in (comp.get("photo_notes") or []) if s(x)]).strip()
            left_text = left_text or "Photos captured during monitoring:"

            for u in photos:
                url = s(u)
                if not url:
                    continue
                b = photo_bytes.get(url)
                if not b:
                    continue

                add_text_left_photo_right_block(
                    doc,
                    left_text=left_text,
                    image_bytes=b,
                    left_width_in=4.2,
                )

        doc.add_paragraph("")
