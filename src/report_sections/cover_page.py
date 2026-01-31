import io
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from PIL import Image
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Mm, Inches


class ReportStyle:
    cover_font: str = "Cambria"
    cover_title_size: int = 18
    cover_table_label_size: int = 11
    cover_table_value_size: int = 10
    title_blue: RGBColor = RGBColor(0, 112, 192)

    header_fill_hex: str = "D9E2F3"
    border_hex: str = "A6A6A6"

    # spacing (tuned to match screenshot)
    title_after_pt: int = 0
    gap_title_to_image_pt: int = 10
    gap_image_to_table_pt: int = 6

    # Cell padding (twips/DXA) - safe for auto-growing rows
    cell_pad_top: int = 40
    cell_pad_bottom: int = 40
    cell_pad_left: int = 110
    cell_pad_right: int = 110

    # REQUIRED image size (inches)
    cover_image_width_in: float = 7.28
    cover_image_height_in: float = 4.84

    # ✅ Table "normal" minimum row height (inches)
    # If text grows -> row grows automatically
    table_row_height_in: float = 0.2


STYLE = ReportStyle()

DATA_KEYS = {
    "PROVINCE": "A01_Province",
    "DISTRICT": "A02_District",
    "VILLAGE": "Village",
    "STARTTIME": "starttime",
    "VISIT_NO": "A26_Visit_number",
    "ACTIVITY_NAME": "Activity_Name",
    "TOOL_NAME": "Tool_Name",
    "PRIMARY_PARTNER": "Primary_Partner_Name",
}


def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def set_run(run, font: str, size: int, bold: bool = False, color: Optional[RGBColor] = None) -> None:
    run.font.name = font
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    except Exception:
        pass
    run.font.size = Pt(int(size))
    run.bold = bool(bold)
    if color is not None:
        run.font.color.rgb = color


def add_spacer(doc, after_pt: int) -> None:
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(int(after_pt))
    p.paragraph_format.line_spacing = 1


# ============================================================
# A4 + margins adjusted slightly so 7.28" fits cleanly
# ============================================================
def set_page_a4(section) -> None:
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # keep narrow-like margins, but allow exact 7.28" image width
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.left_margin = Mm(12.5)
    section.right_margin = Mm(12.5)

    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)


def format_date_dd_mon_yyyy(value: Any) -> str:
    sv = s(value)
    if not sv:
        return ""
    sv_clean = sv.split(".")[0].replace("T", " ").strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            return datetime.strptime(sv_clean, fmt).strftime("%d/%b/%Y")
        except Exception:
            pass
    return sv.split(" ")[0] if " " in sv else sv


def remove_tool_prefix(text: Any) -> str:
    return re.sub(r"^Tool\s*\d+\s*", "", s(text), flags=re.IGNORECASE).strip()


def compose_location(row: Dict[str, Any], overrides: Optional[Dict[str, Any]] = None) -> str:
    overrides = overrides or {}

    def getv(logical_key: str) -> str:
        if logical_key in overrides:
            return s(overrides.get(logical_key))
        ds_key = DATA_KEYS.get(logical_key, "")
        if ds_key and ds_key in overrides:
            return s(overrides.get(ds_key))
        return s(row.get(ds_key))

    parts = [x for x in [getv("PROVINCE"), getv("DISTRICT"), getv("VILLAGE")] if x]
    return ", ".join(parts)


# ============================================================
# Table helpers
# ============================================================
def _emu_to_twips(emu: int) -> int:
    return int(round(int(emu) / 635.0))


def _set_table_fixed_layout(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def _set_table_width_and_indent(table, width_emu: int) -> None:
    width_emu = int(width_emu)
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_emu_to_twips(width_emu)))

    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), "0")


def set_table_borders(table, color_hex: Optional[str] = None) -> None:
    color_hex = (color_hex or STYLE.border_hex).replace("#", "")
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def shade(cell, fill_hex: Optional[str] = None) -> None:
    fill_hex = (fill_hex or STYLE.header_fill_hex).replace("#", "")
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def set_cell_margins(cell, top_dxa=40, bottom_dxa=40, left_dxa=110, right_dxa=110) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _set(side: str, val: int):
        el = tcMar.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")

    _set("top", top_dxa)
    _set("bottom", bottom_dxa)
    _set("start", left_dxa)
    _set("end", right_dxa)


def _tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1


def _cell_write(cell, text: Any, font: str, size: int, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    set_cell_margins(
        cell,
        top_dxa=STYLE.cell_pad_top,
        bottom_dxa=STYLE.cell_pad_bottom,
        left_dxa=STYLE.cell_pad_left,
        right_dxa=STYLE.cell_pad_right,
    )
    p = cell.paragraphs[0]
    _tight_paragraph(p, align)
    r = p.add_run(s(text))
    set_run(r, font, size, bold)


def _row_cant_split(row) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _row_set_min_height(row, height_in: float) -> None:
    """
    ✅ Default row height = height_in
    ✅ If text grows -> row grows automatically
    """
    row.height = Inches(float(height_in))
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


# ============================================================
# Image helpers (NO CROP) -> Fit inside box, pad white
# ============================================================
def _bytes_look_like_html(data: bytes) -> bool:
    if not data:
        return True
    head = data[:400].lower()
    return head.startswith(b"<!doctype html") or b"<html" in head or b"<head" in head


def resize_fit_to_box_png(
    img_bytes: bytes,
    *,
    target_w_in: float,
    target_h_in: float,
    dpi: int = 150,
    pad_color=(255, 255, 255),
) -> bytes:
    if _bytes_look_like_html(img_bytes):
        raise ValueError("Not an image.")

    dpi = max(int(dpi), 72)
    target_w_px = max(int(round(float(target_w_in) * dpi)), 80)
    target_h_px = max(int(round(float(target_h_in) * dpi)), 80)

    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    iw, ih = img.size
    if iw <= 0 or ih <= 0:
        raise ValueError("Invalid image.")

    scale = min(target_w_px / float(iw), target_h_px / float(ih))
    new_w = max(int(round(iw * scale)), 1)
    new_h = max(int(round(ih * scale)), 1)

    img = img.resize((new_w, new_h), Image.LANCZOS)

    canvas = Image.new("RGB", (target_w_px, target_h_px), pad_color)
    x0 = (target_w_px - new_w) // 2
    y0 = (target_h_px - new_h) // 2
    canvas.paste(img, (x0, y0))

    out = io.BytesIO()
    canvas.save(out, format="PNG", optimize=True)
    return out.getvalue()


def add_cover_page(
    doc,
    row: Dict[str, Any],
    cover_image_bytes: Optional[bytes],
    general_info_overrides: Optional[Dict[str, Any]] = None,
) -> None:
    """
    ✅ Without logos (header/footer handled elsewhere)
      - Titles centered blue
      - Cover image EXACT 7.28" × 4.84"
      - Table rows NORMAL min 0.2" and auto-grow if text increases
      - Clean spacing
    """
    ovr = general_info_overrides or {}
    section = doc.sections[0]
    set_page_a4(section)

    # ---------------- Titles ----------------
    title_lines = ("Third-party Monitoring (TPM) — WASH PROGRAMME", "FIELD Visit REPORT")

    for idx, line in enumerate(title_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(STYLE.title_after_pt if idx == 0 else 0)
        p.paragraph_format.line_spacing = 1

        r = p.add_run(line)
        set_run(r, STYLE.cover_font, STYLE.cover_title_size, bold=True, color=STYLE.title_blue)

    add_spacer(doc, STYLE.gap_title_to_image_pt)

    # ---------------- Build table rows ----------------
    def get_value(logical_key: str) -> str:
        if logical_key in ovr:
            return s(ovr.get(logical_key))
        ds = DATA_KEYS.get(logical_key, "")
        if ds and ds in ovr:
            return s(ovr.get(ds))
        return s(row.get(ds)) if ds else ""

    activity_name = get_value("ACTIVITY_NAME")
    visit_no = get_value("VISIT_NO")
    tool_raw = s(ovr.get("TOOL_NAME", row.get(DATA_KEYS["TOOL_NAME"], row.get("Tool", ""))))
    intervention = remove_tool_prefix(tool_raw)
    location = compose_location(row, overrides=ovr)
    visit_date = format_date_dd_mon_yyyy(get_value("STARTTIME"))
    ip_name = get_value("PRIMARY_PARTNER")

    cover_rows: List[Tuple[str, str]] = [
        ("Project Title:", activity_name),
        ("Visit No.:", visit_no),
        ("Type of Intervention:", intervention),
        ("Province / District / Village:", location),
        ("Date of Visit:", visit_date),
        ("Implementing Partner (IP):", ip_name),
        ("Prepared by:", "Premium Performance Consulting (PPC) & Act for Performance"),
        ("Prepared for:", "UNICEF"),
    ]

    # ---------------- Cover image (EXACT SIZE) ----------------
    if cover_image_bytes:
        try:
            clean = resize_fit_to_box_png(
                cover_image_bytes,
                target_w_in=STYLE.cover_image_width_in,
                target_h_in=STYLE.cover_image_height_in,
                dpi=150,
            )
            pimg = doc.add_paragraph()
            pimg.paragraph_format.space_before = Pt(0)
            pimg.paragraph_format.space_after = Pt(0)
            pimg.paragraph_format.line_spacing = 1
            pimg.alignment = WD_ALIGN_PARAGRAPH.CENTER

            pimg.add_run().add_picture(
                io.BytesIO(clean),
                width=Inches(STYLE.cover_image_width_in),
                height=Inches(STYLE.cover_image_height_in),
            )
        except Exception:
            p_bad = doc.add_paragraph("Cover photo could not be embedded.")
            p_bad.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_bad.paragraph_format.space_after = Pt(0)

    add_spacer(doc, STYLE.gap_image_to_table_pt)

    # ---------------- Table ----------------
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table)
    _set_table_fixed_layout(table)

    usable_width_emu = int(section.page_width - section.left_margin - section.right_margin)
    _set_table_width_and_indent(table, usable_width_emu)

    left_w_emu = int(usable_width_emu * 0.32)
    right_w_emu = int(usable_width_emu - left_w_emu)
    table.columns[0].width = left_w_emu
    table.columns[1].width = right_w_emu

    for label, value in cover_rows:
        rr = table.add_row()
        _row_cant_split(rr)
        _row_set_min_height(rr, STYLE.table_row_height_in)  # ✅ min 0.2", auto-grow

        c0, c1 = rr.cells[0], rr.cells[1]
        c0.width = left_w_emu
        c1.width = right_w_emu
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        shade(c0, STYLE.header_fill_hex)
        _cell_write(c0, label, STYLE.cover_font, STYLE.cover_table_label_size, bold=True)
        _cell_write(c1, value, STYLE.cover_font, STYLE.cover_table_value_size, bold=False)
