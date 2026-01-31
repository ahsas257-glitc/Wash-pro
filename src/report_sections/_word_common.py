# src/report_sections/_word_common.py
from __future__ import annotations

import io
import os
import re
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple, Union

from docx.document import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.table import _Row
from docx.text.paragraph import Paragraph


# ============================================================
# Constants
# ============================================================
CHECKED = "☒"
UNCHECKED = "☐"


# ============================================================
# Core text helpers
# ============================================================
def s(v: Any) -> str:
    return "" if v is None else str(v).strip()


def set_run(
    run,
    font: str,
    size: Union[int, float],
    bold: bool = False,
    color: Optional[RGBColor] = None,
) -> None:
    """Safe run styling with eastAsia font set (prevents Arabic/Farsi/Unicode fallback)."""
    try:
        run.font.name = font
    except Exception:
        pass

    # East Asia font (helps prevent fallback rendering issues)
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
    except Exception:
        pass

    try:
        run.font.size = Pt(float(size))
    except Exception:
        pass

    run.bold = bool(bold)

    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass


def tight_paragraph(
    p: Paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before_pt: Union[int, float] = 0,
    after_pt: Union[int, float] = 0,
    line_spacing: float = 1.0,
) -> None:
    """Standard paragraph formatting (tight, predictable)."""
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(float(before_pt))
    pf.space_after = Pt(float(after_pt))
    pf.line_spacing = float(line_spacing)


def add_spacer(doc: Document, *, after_pt: Union[int, float]) -> None:
    p = doc.add_paragraph("")
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1)


def add_body_paragraph(
    doc: Document,
    text: Any,
    *,
    font: str = "Times New Roman",
    size: float = 11,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before_pt: float = 0,
    after_pt: float = 0,
    line_spacing: float = 1.0,
) -> bool:
    """
    Fast common body paragraph writer.
    Returns True if a paragraph was added.
    """
    t = s(text)
    if not t:
        return False

    p = doc.add_paragraph()
    tight_paragraph(p, align=align, before_pt=before_pt, after_pt=after_pt, line_spacing=line_spacing)
    r = p.add_run(t)
    set_run(r, font, size, bold=bold)
    return True


# ============================================================
# Paragraph borders (underline)
# ============================================================
def set_paragraph_bottom_border(
    paragraph: Paragraph,
    *,
    color_hex: str,
    size_eighths: int = 12,
    space: int = 2,
) -> None:
    """Adds/updates a bottom border to a paragraph (idempotent)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)

    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(size_eighths)))
    bottom.set(qn("w:space"), str(int(space)))
    bottom.set(qn("w:color"), s(color_hex).replace("#", ""))


def title_with_orange_line(
    doc: Document,
    *,
    text: str,
    font: str = "Cambria",
    size: Union[int, float] = 14,
    color: RGBColor = RGBColor(0, 112, 192),
    orange_hex: str = "ED7D31",
    after_pt: Union[int, float] = 6,
    border_size_eighths: int = 12,
    border_space: int = 2,
) -> Paragraph:
    """
    Blue title + orange underline line.
    NOTE: Not TOC-safe unless you use add_heading/add_section_title_h1.
    Kept for compatibility (non-TOC titles).
    """
    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1)
    r = p.add_run(s(text))
    set_run(r, font, size, bold=True, color=color)
    set_paragraph_bottom_border(p, color_hex=orange_hex, size_eighths=border_size_eighths, space=border_space)
    return p


# ============================================================
# Outline level (TOC control)
# ============================================================
def set_outline_level(paragraph: Paragraph, level: Optional[int]) -> None:
    """
    level=None -> remove outline level
    level=0 -> Heading 1 outline
    level=1 -> Heading 2 outline ...
    """
    pPr = paragraph._p.get_or_add_pPr()
    outline = pPr.find(qn("w:outlineLvl"))

    if level is None:
        if outline is not None:
            pPr.remove(outline)
        return

    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        pPr.append(outline)

    outline.set(qn("w:val"), str(int(level)))


# ============================================================
# TOC-safe Heading helpers
# ============================================================
def add_heading(
    doc: Document,
    text: str,
    *,
    level: int = 1,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font: str = "Cambria",
    size: int = 14,
    bold: bool = True,
    color: Optional[RGBColor] = None,
) -> Paragraph:
    """
    Add a REAL Word Heading (Heading 1/2/3/...)
    REQUIRED for Table of Contents to work correctly.
    """
    lvl = max(1, min(int(level), 9))

    p = doc.add_paragraph(s(text))
    try:
        p.style = f"Heading {lvl}"
    except Exception:
        pass

    p.alignment = align

    # style runs
    for r in p.runs:
        r.bold = bool(bold)
        try:
            r.font.name = font
        except Exception:
            pass
        try:
            r._element.rPr.rFonts.set(qn("w:eastAsia"), font)  # type: ignore[attr-defined]
        except Exception:
            pass
        try:
            r.font.size = Pt(int(size))
        except Exception:
            pass
        if color is not None:
            try:
                r.font.color.rgb = color
            except Exception:
                pass

    return p


def add_section_title_h1(
    doc: Document,
    text: str,
    *,
    font: str = "Cambria",
    size: Union[int, float] = 16,
    color: RGBColor = RGBColor(0, 112, 192),
    orange_hex: str = "ED7D31",
    after_pt: Union[int, float] = 6,
    border_size_eighths: int = 12,
    border_space: int = 2,
) -> Paragraph:
    """
    Standard section title:
      - Real Heading 1 (TOC-friendly)
      - Font size 16
      - Blue title + orange underline (on SAME paragraph)
    """
    p = add_heading(
        doc,
        s(text),
        level=1,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        font=font,
        size=int(size),
        bold=True,
        color=color,
    )
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=float(after_pt), line_spacing=1)
    set_paragraph_bottom_border(
        p,
        color_hex=orange_hex,
        size_eighths=int(border_size_eighths),
        space=int(border_space),
    )
    return p


# ============================================================
# Word field-code helpers (TOC, PAGEREF, etc.)
# ============================================================
def _add_field_run(paragraph: Paragraph, field_instr: str) -> None:
    """Adds a Word field to a paragraph (Word renders it when fields are updated)."""
    # begin
    r_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(fld_begin)

    # instruction text
    r_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_instr
    r_instr._r.append(instr)

    # separate
    r_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    r_sep._r.append(fld_sep)

    # placeholder result
    paragraph.add_run(" ")

    # end
    r_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fld_end)


def add_toc_field(
    doc: Document,
    *,
    levels: str = "1-3",
    hyperlinks: bool = True,
    hide_page_numbers_in_web_layout: bool = False,
) -> Paragraph:
    """Insert a TOC field in the document (Word updates fields on open / update)."""
    switches = [rf'\\o "{levels}"', r"\u"]
    if hyperlinks:
        switches.append(r"\h")
    if hide_page_numbers_in_web_layout:
        switches.append(r"\z")

    instr = "TOC " + " ".join(switches)
    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
    _add_field_run(p, instr)
    return p


def update_docx_fields_bytes(docx_bytes: bytes) -> bytes:
    """
    Best-effort update of Word fields (TOC, page numbers).
    On Windows with MS Word installed, this will fully update fields.
    Otherwise returns original bytes.
    """
    if not docx_bytes:
        return docx_bytes

    try:
        import win32com.client  # requires pywin32 + MS Word
    except Exception:
        return docx_bytes

    with tempfile.TemporaryDirectory() as td:
        path = os.path.join(td, "report.docx")
        with open(path, "wb") as f:
            f.write(docx_bytes)

        word = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            wdoc = word.Documents.Open(path)
            wdoc.Fields.Update()

            # Update TOC objects explicitly
            try:
                for i in range(1, wdoc.TablesOfContents.Count + 1):
                    wdoc.TablesOfContents(i).Update()
            except Exception:
                pass

            wdoc.Repaginate()
            wdoc.Save()
            wdoc.Close()
            word.Quit()

            with open(path, "rb") as f:
                return f.read()

        except Exception:
            try:
                if word is not None:
                    word.Quit()
            except Exception:
                pass
            return docx_bytes


# ============================================================
# Page helpers (A4 + usable width)
# ============================================================
def set_page_a4(
    section,
    *,
    page_w_mm: int = 210,
    page_h_mm: int = 297,
    margin_top_mm: int = 8,
    margin_bottom_mm: int = 8,
    margin_left_mm: int = 12,
    margin_right_mm: int = 12,
    header_dist_mm: int = 5,
    footer_dist_mm: int = 5,
) -> None:
    section.page_width = Mm(int(page_w_mm))
    section.page_height = Mm(int(page_h_mm))
    section.top_margin = Mm(int(margin_top_mm))
    section.bottom_margin = Mm(int(margin_bottom_mm))
    section.left_margin = Mm(int(margin_left_mm))
    section.right_margin = Mm(int(margin_right_mm))
    section.header_distance = Mm(int(header_dist_mm))
    section.footer_distance = Mm(int(footer_dist_mm))


def emu_to_mm(emu: int) -> int:
    # 1 mm = 36000 EMU
    return int(round(int(emu) / 36000.0))


def section_usable_width_emu(section) -> int:
    return int(section.page_width.emu - section.left_margin.emu - section.right_margin.emu)


def section_usable_width_mm(section) -> int:
    return emu_to_mm(section_usable_width_emu(section))


# ============================================================
# Table helpers (layout, borders, shading, padding)
# ============================================================
def _emu_to_twips(emu: int) -> int:
    # 1 inch = 914400 EMU = 1440 twips => twips = emu / 635
    return int(round(int(emu) / 635.0))


def set_table_fixed_layout(table) -> None:
    """Forces table layout to fixed (prevents Word auto-resizing columns)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")


def set_table_width_exact(table, width) -> None:
    """Sets exact table width. width can be docx Length OR integer EMU."""
    width_emu = int(width.emu) if hasattr(width, "emu") else int(width)

    tbl = table._tbl
    tblPr = tbl.tblPr

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(_emu_to_twips(width_emu)))

    # remove indentation to keep it aligned
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:type"), "dxa")
    tblInd.set(qn("w:w"), "0")


def set_table_width_from_section(table, doc: Document, *, section_index: int = 0) -> None:
    """Sets table width to usable content width of a given section."""
    idx = max(0, min(int(section_index), len(doc.sections) - 1))
    sec = doc.sections[idx]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    set_table_width_exact(table, int(usable))


def set_table_columns_exact(table, widths_in: List[float]) -> None:
    """
    Force exact column widths by writing:
      - w:tblGrid / w:gridCol
      - w:tcW for every cell
    This prevents Word from reverting column widths.
    widths_in: list of inches, e.g. [2.30, 4.60]
    """
    def _in_to_tw(inches: float) -> int:
        return int(round(float(inches) * 1440))  # 1 inch = 1440 twips

    widths_tw = [_in_to_tw(w) for w in (widths_in or []) if float(w) > 0]
    if not widths_tw:
        return

    tbl = table._tbl
    tblPr = tbl.tblPr

    # ensure fixed layout
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    # tblGrid
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        # insert after tblPr if possible
        try:
            tbl.insert(1, tblGrid)
        except Exception:
            tbl.append(tblGrid)

    # clear existing grid cols
    for gc in list(tblGrid):
        tblGrid.remove(gc)

    for w in widths_tw:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w)))
        tblGrid.append(gc)

    # apply tcW to all existing cells
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            w = widths_tw[min(i, len(widths_tw) - 1)]
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(int(w)))


def set_table_borders(table, *, color_hex: str = "000000", size: str = "8") -> None:
    """Applies single-line borders to the whole table (idempotent)."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), s(color_hex).replace("#", ""))


def set_table_borders_none(table) -> None:
    """Removes borders from a table (idempotent). Great for layout tables."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def shade_cell(cell, fill_hex: str) -> None:
    """Cell background shading (idempotent)."""
    h = s(fill_hex).replace("#", "").upper()
    if not h:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), h)


def set_cell_margins(
    cell,
    top_dxa: int = 120,
    bottom_dxa: int = 120,
    left_dxa: int = 140,
    right_dxa: int = 140,
) -> None:
    """Cell padding in DXA (twips). Uses start/end (RTL friendly)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    def _set(side: str, val: int) -> None:
        tag = f"w:{side}"
        el = tcMar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")

    _set("top", top_dxa)
    _set("bottom", bottom_dxa)
    _set("start", left_dxa)
    _set("end", right_dxa)


def set_cell_borders(cell, *, size: int = 8, color_hex: str = "000000") -> None:
    """Adds borders for a cell (idempotent). size in 1/8 pt units."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(int(size)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), s(color_hex).replace("#", ""))


def set_cell_borders_none(cell) -> None:
    """Removes borders for a cell (idempotent)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "nil")


def set_row_cant_split(row: Union[_Row, object], *, cant_split: bool = True) -> None:
    """Prevent row from splitting across pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    el = trPr.find(qn("w:cantSplit"))
    if cant_split:
        if el is None:
            trPr.append(OxmlElement("w:cantSplit"))
    else:
        if el is not None:
            trPr.remove(el)


def _inches_to_twips(inches: float) -> int:
    return int(round(float(inches) * 1440))


def set_row_height_exact(row, height_in: float) -> None:
    """Sets exact row height in inches (may clip if content exceeds)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)

    trHeight.set(qn("w:val"), str(_inches_to_twips(height_in)))
    trHeight.set(qn("w:hRule"), "exact")


def set_row_height_at_least(row, height_in: float) -> None:
    """Sets minimum row height in inches (can grow)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()

    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)

    trHeight.set(qn("w:val"), str(_inches_to_twips(height_in)))
    trHeight.set(qn("w:hRule"), "atLeast")


def set_repeat_table_header(row) -> None:
    """Make a table row repeat as header on each new page in Word."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = trPr.find(qn("w:tblHeader"))
    if tblHeader is None:
        tblHeader = OxmlElement("w:tblHeader")
        trPr.append(tblHeader)
    tblHeader.set(qn("w:val"), "true")


# ============================================================
# Heading numbering strip + consistent cell writer
# ============================================================
_NUM_PREFIX_RE = re.compile(r"^\s*\d+(?:\.\d+)*\s*[\.\-\)]?\s*", re.UNICODE)

def strip_heading_numbering(title: str) -> str:
    """'5.1. X' -> 'X'"""
    return _NUM_PREFIX_RE.sub("", (title or "").strip()).strip()


def write_cell_text(
    cell,
    text: str,
    *,
    font: str = "Times New Roman",
    size: Union[int, float] = 11,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Clear cell and write a single paragraph with standard formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1)
    r = p.add_run(text or "")
    set_run(r, font, size, bold=bold)


# ============================================================
# Checkbox + bullets helpers
# ============================================================
_BULLET_SPLIT_RE = re.compile(r"(?:\r?\n)+", re.UNICODE)

def parse_bool_like(v: Any) -> Optional[bool]:
    """
    Returns True/False/None.
    Accepts: bool, 0/1, yes/no, checked/unchecked, ✅/❌, etc.
    """
    if v is None:
        return None
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        iv = int(v)
        if iv == 1:
            return True
        if iv == 0:
            return False

    sv = s(v).strip().lower()
    if sv in {"yes", "y", "true", "1", "checked", "✔", "✅"}:
        return True
    if sv in {"no", "n", "false", "0", "unchecked", "✘", "❌"}:
        return False
    return None


def as_yes(v: Any) -> bool:
    return s(v).lower() in {"yes", "y", "true", "1", "checked", "✔", "✅"}


def as_no(v: Any) -> bool:
    return s(v).lower() in {"no", "n", "false", "0", "unchecked", "✘", "❌"}


def checkbox_line(value: Optional[bool], *, yes_label: str = "Yes", no_label: str = "No") -> str:
    """Creates: ☒ Yes   ☐ No  (None -> both unchecked)."""
    yes = f"{CHECKED} {yes_label}" if value is True else f"{UNCHECKED} {yes_label}"
    no = f"{CHECKED} {no_label}" if value is False else f"{UNCHECKED} {no_label}"
    return f"{yes}   {no}"


def bullets_from_text(text: Any) -> List[str]:
    """Multiline text -> bullet items (trim leading '-', '•')."""
    raw = s(text)
    if not raw:
        return []
    parts = [p.strip() for p in _BULLET_SPLIT_RE.split(raw) if p.strip()]
    out: List[str] = []
    for p in parts:
        p = p.lstrip("•- \t").strip()
        if p:
            out.append(p)
    return out


def add_bullets_paragraphs(
    cell,
    items: List[str],
    *,
    font: str = "Times New Roman",
    size: Union[int, float] = 11,
) -> None:
    """
    Adds bullet-like paragraphs inside a DOCX cell using plain text bullets.
    Faster than numbering API and stable in Word.
    """
    if not items:
        return

    cell.text = ""
    for it in items:
        p = cell.add_paragraph()
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        r = p.add_run(f"• {it}")
        set_run(r, font, size, bold=False)


def write_yes_no_checkboxes(
    cell,
    value: Any,
    *,
    font_size: Union[int, float] = 11,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    b = parse_bool_like(value)
    yes = f"{CHECKED} Yes" if b is True else f"{UNCHECKED} Yes"
    no = f"{CHECKED} No" if b is False else f"{UNCHECKED} No"

    cell.text = ""
    p = cell.paragraphs[0]
    tight_paragraph(p, align=align, before_pt=0, after_pt=0, line_spacing=1)
    set_run(p.add_run(f"{yes}  {no}"), "Times New Roman", font_size, False)


def write_two_option_checkboxes(
    cell,
    value: Any,
    opt1: str,
    opt2: str,
    *,
    font_size: Union[int, float] = 11,
) -> None:
    checked_opt1 = False
    checked_opt2 = False

    if isinstance(value, (int, float)):
        iv = int(value)
        checked_opt1 = (iv == 1)
        checked_opt2 = (iv == 0)
    else:
        v = s(value).lower()
        if v in {"1", "male", "m", "man"}:
            checked_opt1 = True
        elif v in {"0", "female", "f", "woman"}:
            checked_opt2 = True

    a = f"{CHECKED} {opt1}" if checked_opt1 else f"{UNCHECKED} {opt1}"
    b = f"{CHECKED} {opt2}" if checked_opt2 else f"{UNCHECKED} {opt2}"

    cell.text = ""
    p = cell.paragraphs[0]
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
    set_run(p.add_run(f"{a}  {b}"), "Times New Roman", font_size, False)


def three_option_checkbox_line(value: Any, a: str, b: str, c: str) -> str:
    v = s(value).lower()
    aa, bb, cc = a.lower(), b.lower(), c.lower()

    def box(txt: str, on: bool) -> str:
        return f"{CHECKED if on else UNCHECKED} {txt}"

    return "    ".join([box(a, v == aa), box(b, v == bb), box(c, v == cc)])


# ============================================================
# Executive-summary shared helpers
# ============================================================
def pick_first_nonempty(row: Dict[str, Any], overrides: Dict[str, Any], keys: List[str]) -> Any:
    """Pick first non-empty value from overrides, else from row."""
    overrides = overrides or {}
    row = row or {}

    for k in keys:
        v = overrides.get(k, None)
        if v is not None and s(v):
            return v
    for k in keys:
        v = row.get(k, None)
        if v is not None and s(v):
            return v
    return None


def date_only_isoish(v: Any) -> str:
    """
    Accepts:
      - '2026-01-30T12:00:00'
      - '2026-01-30 12:00:00'
      - '2026-01-30'
    Returns: '2026-01-30'
    """
    t = s(v)
    if not t:
        return ""
    t = t.replace("T", " ").strip()
    return t.split(" ")[0].strip()


def norm_phrase(v: Any) -> str:
    """underscores -> spaces, lowercase, trim."""
    return s(v).lower().replace("_", " ").strip()


def build_location_phrase(village: str, district: str, province: str) -> str:
    parts: List[str] = []
    v = s(village)
    d = s(district)
    p = s(province)

    if v:
        parts.append(v)
    if d:
        parts.append(f"{d} District")
    if p:
        parts.append(f"{p} Province")
    return ", ".join(parts)


# ============================================================
# General Project Information helpers (shared)
# ============================================================
_DIGITS_RE = re.compile(r"\D+", re.UNICODE)

def na_if_empty(v: Any) -> str:
    t = s(v)
    return t if t else "N/A"


def format_af_phone(raw: Any) -> str:
    txt = _DIGITS_RE.sub("", s(raw))
    if not txt:
        return ""
    if txt.startswith("0"):
        txt = txt[1:]
    if txt.startswith("93"):
        return f"+{txt}"
    return f"+93{txt}"


def normalize_email_or_na_strict(raw: Any) -> str:
    t = s(raw)
    if not t or "@" not in t or "." not in t:
        return "N/A"
    return t


def donor_upper_and_pipe(v: Any) -> str:
    t = s(v)
    return t.upper().replace(",", " | ") if t else ""


def extract_date_only(value: Any) -> str:
    sv = s(value)
    if not sv:
        return ""
    return sv.split(" ")[0].strip() if " " in sv else sv


def format_date_dd_mon_yyyy(value: Any) -> str:
    sv = s(value)
    if not sv:
        return ""
    sv = sv.replace("T", " ").split(".")[0].strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(sv, fmt)
            return dt.strftime("%d/%b/%Y")
        except Exception:
            pass
    return extract_date_only(sv)


def add_available_documents_inner_table(details_cell, row: Dict[str, Any], keys: Dict[str, str]) -> None:
    """
    Nested 4x4 table used in General Project Information:
      - left label | left checkbox | right label | right checkbox
    """
    details_cell.text = ""
    set_cell_margins(details_cell, 0, 0, 0, 0)

    inner = details_cell.add_table(rows=4, cols=4)
    inner.autofit = False
    inner.style = "Table Grid"
    set_table_fixed_layout(inner)

    col_widths = [Inches(x) for x in (0.71, 1.10, 1.82, 1.10)]
    for i, w in enumerate(col_widths):
        inner.columns[i].width = w

    set_table_borders(inner, color_hex="FFFFFF")

    left_items = [
        ("Contract", row.get(keys["DOC_CONTRACT"])),
        ("Journal", row.get(keys["DOC_JOURNAL"])),
        ("BOQ", row.get(keys["DOC_BOQ"])),
        ("Drawing", row.get(keys["DOC_DRAWINGS"])),
    ]
    right_items = [
        ("Site engineer", row.get(keys["DOC_SITE_ENGINEER"])),
        ("geophysical, hydrological tests", row.get(keys["DOC_GEOPHYSICAL"])),
        ("water quality tests", row.get(keys["DOC_WQ_TEST"])),
        ("Pump test result", row.get(keys["DOC_PUMP_TEST"])),
    ]

    for i in range(4):
        c00, c01, c02, c03 = inner.cell(i, 0), inner.cell(i, 1), inner.cell(i, 2), inner.cell(i, 3)

        for c in (c00, c01, c02, c03):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(c, "F2F2F2")
            set_cell_margins(c, 60, 60, 90, 90)

        write_cell_text(c00, left_items[i][0], font="Times New Roman", size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_yes_no_checkboxes(c01, left_items[i][1], font_size=11, align=WD_ALIGN_PARAGRAPH.LEFT)

        write_cell_text(c02, right_items[i][0], font="Times New Roman", size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_yes_no_checkboxes(c03, right_items[i][1], font_size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)


# ============================================================
# Text normalization + recommendation splitting
# ============================================================
_WS_RE = re.compile(r"\s+", re.UNICODE)
_MULTI_DOT_RE = re.compile(r"\.\s*\.+", re.UNICODE)
_RECO_SPLIT_RE = re.compile(r";\s+|\.\s+(?=[A-Z])", re.UNICODE)

def normalize_sentence(text: Any) -> str:
    t = s(text)
    if not t:
        return ""
    t = _WS_RE.sub(" ", t).strip()
    t = _MULTI_DOT_RE.sub(".", t)
    return t


def split_recommendations_conservative(text: Any) -> List[str]:
    """
    Conservative split:
      - if multiple lines -> each line is an item
      - else split by ';' or '. ' (only if next looks like sentence start)
    """
    raw = s(text)
    t = normalize_sentence(raw)
    if not t:
        return []

    if "\n" in raw:
        parts = [p.strip(" -•\t").strip() for p in raw.split("\n") if p.strip()]
        return [p for p in parts if p]

    parts = [p.strip() for p in _RECO_SPLIT_RE.split(t) if p.strip()]
    return parts if parts else [t]


def severity_checkbox_line(selected: str = "") -> str:
    """
    ☐ Low  ☐ Medium  ☐ High
    If selected matches one -> checked.
    """
    sel = s(selected).strip().lower()

    def box(lbl: str) -> str:
        l = lbl.lower()
        on = (sel == l) or (sel in {"h", "m", "l"} and sel == l[0])
        return f"{CHECKED if on else UNCHECKED} {lbl}"

    return f"{box('Low')}    {box('Medium')}    {box('High')}"


# ============================================================
# Findings/Recommendations extraction helpers
# ============================================================
def extract_findings_and_recos_from_section5(component_observations: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    """
    Output: [{"finding": "...", "recommendation": "..."}]
    """
    rows: List[Dict[str, str]] = []

    for comp in (component_observations or []):
        if not isinstance(comp, dict):
            continue

        findings: List[str] = []
        for sub in (comp.get("subsections") or []):
            mt = sub.get("major_table")
            if isinstance(mt, list):
                for rr in mt:
                    if not isinstance(rr, dict):
                        continue
                    ftxt = normalize_sentence(rr.get("Findings") or rr.get("Finding") or rr.get("finding"))
                    if ftxt:
                        findings.append(ftxt)

        if not findings:
            continue

        reco_texts: List[str] = []
        for sub in (comp.get("subsections") or []):
            st = s(sub.get("title")).lower()
            if "recommendation" in st:
                for para in (sub.get("paragraphs") or []):
                    if s(para):
                        reco_texts.append(s(para))

        reco_joined = normalize_sentence("\n".join(reco_texts))
        reco_items = split_recommendations_conservative(reco_joined)

        if reco_items and len(reco_items) == len(findings):
            for f, r in zip(findings, reco_items):
                rows.append({"finding": f, "recommendation": normalize_sentence(r)})
        else:
            for f in findings:
                rows.append({"finding": f, "recommendation": reco_joined})

    cleaned: List[Dict[str, str]] = []
    for rr in rows:
        f = normalize_sentence(rr.get("finding", ""))
        r = normalize_sentence(rr.get("recommendation", ""))
        if f:
            cleaned.append({"finding": f, "recommendation": r})
    return cleaned


def present_severities_from_mapping(
    extracted_rows: List[Dict[str, str]],
    severity_by_no: Dict[int, str],
    severity_by_finding: Dict[str, str],
) -> List[str]:
    present = set()

    for idx, rr in enumerate(extracted_rows, start=1):
        chosen = ""
        if idx in (severity_by_no or {}):
            chosen = s((severity_by_no or {}).get(idx))
        else:
            ftxt = normalize_sentence(rr.get("finding", ""))
            if ftxt:
                for k, v in (severity_by_finding or {}).items():
                    if normalize_sentence(k).lower() == ftxt.lower():
                        chosen = s(v)
                        break

        low = chosen.strip().lower()
        if low in {"high", "h"}:
            present.add("High")
        elif low in {"medium", "med", "m"}:
            present.add("Medium")
        elif low in {"low", "l"}:
            present.add("Low")

    order = ["High", "Medium", "Low"]
    return [x for x in order if x in present]


# ============================================================
# Section 5 helpers (Component-wise key observations)
# ============================================================
_SPECIAL_TITLE_RE = re.compile(r"[^a-z0-9]+", re.IGNORECASE)

def special_title_normalized(title: str) -> str:
    t = s(title).strip().lower()
    t = t.rstrip(":").strip()
    t = _SPECIAL_TITLE_RE.sub(" ", t)
    return " ".join(t.split())


# ============================================================
# Image helpers (DOCX-safe + FIT/EXACT box)
# ============================================================
def bytes_look_like_html(data: bytes) -> bool:
    if not data:
        return True
    head = data[:400].lower()
    return head.startswith(b"<!doctype html") or b"<html" in head or b"<head" in head


def normalize_image_bytes_for_docx(img_bytes: bytes) -> bytes:
    """Converts input bytes to a safe PNG for docx (rejects HTML/login)."""
    if bytes_look_like_html(img_bytes):
        raise ValueError("Not an image (HTML/login page).")

    from PIL import Image  # lazy import
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PNG", optimize=True)
    return out.getvalue()


def normalize_image_bytes_for_docx_light(image_bytes: bytes, *, max_px: int = 1600) -> bytes:
    """
    Fast path:
      - if already jpg/png -> return as-is
      - else try PIL convert to JPEG (smaller) + resize
    """
    if not image_bytes:
        return b""

    head = image_bytes[:10]
    is_jpeg = head.startswith(b"\xff\xd8\xff")
    is_png = head.startswith(b"\x89PNG\r\n\x1a\n")
    if is_jpeg or is_png:
        return image_bytes

    try:
        from PIL import Image  # lazy import
        im = Image.open(io.BytesIO(image_bytes))
        im.load()

        w, h = im.size
        m = max(w, h)
        if m > max_px:
            scale = max_px / float(m)
            im = im.resize((int(w * scale), int(h * scale)))

        if im.mode != "RGB":
            im = im.convert("RGB")

        out = io.BytesIO()
        im.save(out, format="JPEG", quality=85, optimize=True)
        return out.getvalue()
    except Exception:
        return image_bytes


def resize_crop_to_box_png(
    img_bytes: bytes,
    *,
    target_w_mm: int,
    target_h_mm: int,
    dpi: int = 150,
    optimize: bool = True,
) -> bytes:
    """EXACT box WITH CROP: center crop to aspect ratio, then resize."""
    if bytes_look_like_html(img_bytes):
        raise ValueError("Not an image (HTML/login page).")

    from PIL import Image  # lazy import

    dpi = max(int(dpi), 72)
    tw_in = float(target_w_mm) / 25.4
    th_in = float(target_h_mm) / 25.4
    target_w_px = max(int(round(tw_in * dpi)), 80)
    target_h_px = max(int(round(th_in * dpi)), 80)

    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    iw, ih = img.size
    if iw <= 0 or ih <= 0:
        raise ValueError("Invalid image.")

    target_ratio = target_w_px / float(target_h_px)
    src_ratio = iw / float(ih)

    if src_ratio > target_ratio:
        new_w = int(round(ih * target_ratio))
        x0 = max((iw - new_w) // 2, 0)
        img = img.crop((x0, 0, x0 + new_w, ih))
    else:
        new_h = int(round(iw / target_ratio))
        y0 = max((ih - new_h) // 2, 0)
        img = img.crop((0, y0, iw, y0 + new_h))

    img = img.resize((target_w_px, target_h_px), Image.LANCZOS)
    out = io.BytesIO()
    img.save(out, format="PNG", optimize=bool(optimize))
    return out.getvalue()


def resize_fit_to_box_png(
    img_bytes: bytes,
    *,
    target_w_mm: int,
    target_h_mm: int,
    dpi: int = 150,
    optimize: bool = True,
    pad_color=(255, 255, 255),
) -> bytes:
    """
    NO CROP (fit):
      - keeps full image visible
      - resizes to fit inside target box
      - pads remaining area (white)
    """
    if bytes_look_like_html(img_bytes):
        raise ValueError("Not an image (HTML/login page).")

    from PIL import Image  # lazy import

    dpi = max(int(dpi), 72)
    tw_in = float(target_w_mm) / 25.4
    th_in = float(target_h_mm) / 25.4
    target_w_px = max(int(round(tw_in * dpi)), 80)
    target_h_px = max(int(round(th_in * dpi)), 80)

    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    iw, ih = img.size
    if iw <= 0 or ih <= 0:
        raise ValueError("Invalid image.")

    scale = min(target_w_px / float(iw), target_h_px / float(ih))
    new_w = max(int(round(iw * scale)), 1)
    new_h = max(int(round(ih * scale)), 1)

    img = img.resize((new_w, new_h), Image.LANCZOS)

    canvas = Image.new("RGB", (target_w_px, target_h_px), pad_color)
    x0 = (target_w_px - new_w) // 2
    y0 = (target_h_px - new_h) // 2
    canvas.paste(img, (x0, y0))

    out = io.BytesIO()
    canvas.save(out, format="PNG", optimize=bool(optimize))
    return out.getvalue()


def add_picture_in_cell_exact(
    cell,
    img_bytes: bytes,
    *,
    target_w_mm: int,
    target_h_mm: int,
    dpi: int = 150,
    optimize: bool = True,
) -> bool:
    """Insert exact-sized image in a cell (center crop + resize)."""
    try:
        clean = resize_crop_to_box_png(
            img_bytes,
            target_w_mm=int(target_w_mm),
            target_h_mm=int(target_h_mm),
            dpi=int(dpi),
            optimize=bool(optimize),
        )

        cell.text = ""
        set_cell_margins(cell, top_dxa=0, bottom_dxa=0, left_dxa=0, right_dxa=0)

        p = cell.paragraphs[0]
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)

        run = p.add_run()
        run.add_picture(io.BytesIO(clean), width=Mm(int(target_w_mm)), height=Mm(int(target_h_mm)))
        return True
    except Exception:
        return False


def add_picture_in_cell_fit(
    cell,
    img_bytes: bytes,
    *,
    target_w_mm: int,
    target_h_mm: int,
    dpi: int = 150,
    optimize: bool = True,
) -> bool:
    """Insert image in a fixed box WITHOUT CROPPING (fit + padding)."""
    try:
        clean = resize_fit_to_box_png(
            img_bytes,
            target_w_mm=int(target_w_mm),
            target_h_mm=int(target_h_mm),
            dpi=int(dpi),
            optimize=bool(optimize),
        )

        cell.text = ""
        set_cell_margins(cell, top_dxa=0, bottom_dxa=0, left_dxa=0, right_dxa=0)

        p = cell.paragraphs[0]
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)

        run = p.add_run()
        run.add_picture(io.BytesIO(clean), width=Mm(int(target_w_mm)), height=Mm(int(target_h_mm)))
        return True
    except Exception:
        return False


# ============================================================
# Logo helpers
# ============================================================
def add_center_logo(doc: Document, *, logo_path: str, width_mm: int, after_pt: int = 2) -> None:
    if not logo_path:
        return
    try:
        p = doc.add_paragraph()
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=int(after_pt), line_spacing=1)
        run = p.add_run()
        run.add_picture(logo_path, width=Mm(int(width_mm)))
    except Exception:
        pass


def add_bottom_logos(
    doc: Document,
    *,
    left_logo_path: str,
    right_logo_path: str,
    height_mm: int = 14,
    gap_before_pt: int = 10,
) -> None:
    if gap_before_pt:
        add_spacer(doc, after_pt=int(gap_before_pt))

    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    set_table_fixed_layout(t)
    set_table_borders_none(t)

    c0, c1 = t.cell(0, 0), t.cell(0, 1)
    set_cell_margins(c0, 0, 0, 0, 0)
    set_cell_margins(c1, 0, 0, 0, 0)

    c0.text = ""
    p0 = c0.paragraphs[0]
    tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
    if left_logo_path:
        try:
            r0 = p0.add_run()
            r0.add_picture(left_logo_path, height=Mm(int(height_mm)))
        except Exception:
            pass

    c1.text = ""
    p1 = c1.paragraphs[0]
    tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.RIGHT, before_pt=0, after_pt=0, line_spacing=1)
    if right_logo_path:
        try:
            r1 = p1.add_run()
            r1.add_picture(right_logo_path, height=Mm(int(height_mm)))
        except Exception:
            pass


# ============================================================
# Section 5 table/photo blocks (used by component_wise_key_observations)
# ============================================================
def add_major_findings_table_tool6(
    doc: Document,
    *,
    major_rows: List[Dict[str, Any]],
    photo_bytes: Dict[str, bytes],
    photo_field_map: Optional[Dict[str, str]] = None,
) -> None:
    """
    Table columns: NO | Findings | Compliance | Photo
    """
    photo_field_map = photo_field_map or {}
    major_rows = major_rows or []

    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    tbl.style = "Table Grid"
    set_table_fixed_layout(tbl)

    widths = [Inches(0.45), Inches(3.60), Inches(0.95), Inches(2.00)]
    headers = ["NO", "Findings", "Compliance", "Photo"]

    r0 = tbl.rows[0]
    set_row_cant_split(r0, cant_split=True)
    for i, c in enumerate(r0.cells):
        c.width = widths[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(c, "D9E2F3")
        set_cell_margins(c, 80, 80, 120, 120)

        c.text = ""
        p = c.paragraphs[0]
        tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p.add_run(headers[i]), "Times New Roman", 10, True)

    for rr in major_rows:
        row = tbl.add_row()
        set_row_cant_split(row, cant_split=True)

        cells = row.cells
        for i, c in enumerate(cells):
            c.width = widths[i]
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(c, 80, 80, 120, 120)
            c.text = ""

        no = s(rr.get("NO") or rr.get("No") or rr.get("no"))
        finding = s(rr.get("Findings") or rr.get("Finding") or rr.get("finding"))
        compliance = s(rr.get("Compliance") or rr.get("compliance"))
        photo_url = s(rr.get("Photo") or rr.get("photo"))

        p0 = cells[0].paragraphs[0]
        tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p0.add_run(no), "Times New Roman", 10, False)

        p1 = cells[1].paragraphs[0]
        tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before_pt=0, after_pt=0, line_spacing=1.15)
        set_run(p1.add_run(finding), "Times New Roman", 10, False)

        p2 = cells[2].paragraphs[0]
        tight_paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p2.add_run(compliance), "Times New Roman", 10, False)

        p3 = cells[3].paragraphs[0]
        tight_paragraph(p3, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)

        if photo_url and photo_url in (photo_bytes or {}):
            try:
                clean = normalize_image_bytes_for_docx_light(photo_bytes[photo_url])
                run = p3.add_run()
                run.add_picture(io.BytesIO(clean), width=Inches(1.80))
            except Exception:
                lbl = photo_field_map.get(photo_url, "Photo")
                set_run(p3.add_run(lbl), "Times New Roman", 9, False)
        else:
            lbl = photo_field_map.get(photo_url, "") if photo_url else ""
            if lbl:
                set_run(p3.add_run(lbl), "Times New Roman", 9, False)


def add_text_left_photo_right_block(
    doc: Document,
    *,
    left_text: str,
    image_bytes: bytes,
    left_width_in: float = 4.2,
) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    set_table_fixed_layout(tbl)

    row = tbl.rows[0]
    set_row_cant_split(row, cant_split=True)

    cL, cR = row.cells[0], row.cells[1]
    cL.width = Inches(float(left_width_in))
    cR.width = Inches(2.0)

    set_cell_borders_none(cL)
    set_cell_borders_none(cR)

    cL.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cR.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cL.text = ""
    pL = cL.paragraphs[0]
    tight_paragraph(pL, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before_pt=0, after_pt=0, line_spacing=1.15)
    set_run(pL.add_run(s(left_text)), "Times New Roman", 11, False)

    cR.text = ""
    pR = cR.paragraphs[0]
    tight_paragraph(pR, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=0, line_spacing=1)

    clean = normalize_image_bytes_for_docx_light(image_bytes)
    run = pR.add_run()
    run.add_picture(io.BytesIO(clean), width=Inches(2.0))

    doc.add_paragraph("")


# ============================================================
# Simple body helpers (kept for compatibility)
# ============================================================
def body(doc: Document, text: str) -> None:
    add_body_paragraph(doc, text, font="Times New Roman", size=11, bold=False, after_pt=3, line_spacing=1.15)


def normal_subtitle(doc: Document, text: str) -> None:
    t = s(text)
    if not t:
        return
    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=6, after_pt=3, line_spacing=1)
    set_run(p.add_run(t), "Times New Roman", 11, True)


# ============================================================
# Cover page helpers
# ============================================================
def cover_set_page_margins(doc: Document, *, top_mm: float, bottom_mm: float, left_mm: float, right_mm: float) -> None:
    sec = doc.sections[0]
    sec.top_margin = Mm(top_mm)
    sec.bottom_margin = Mm(bottom_mm)
    sec.left_margin = Mm(left_mm)
    sec.right_margin = Mm(right_mm)


def cover_add_logo_centered(doc: Document, image_bytes: bytes, *, width_in: float = 1.6, space_after_pt: int = 6) -> None:
    if not image_bytes:
        return
    try:
        clean = normalize_image_bytes_for_docx(image_bytes)
    except Exception:
        clean = image_bytes

    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=space_after_pt, line_spacing=1)
    run = p.add_run()
    run.add_picture(io.BytesIO(clean), width=Inches(float(width_in)))


def cover_add_center_title_block(
    doc: Document,
    *,
    title: str,
    subtitle: str = "",
    emphasis: str = "",
) -> None:
    p1 = doc.add_paragraph()
    tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=6, after_pt=6, line_spacing=1)
    r1 = p1.add_run(s(title))
    set_run(r1, "Cambria", 18, bold=True, color=RGBColor(0, 112, 192))

    if s(subtitle):
        p2 = doc.add_paragraph()
        tight_paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=8, line_spacing=1)
        r2 = p2.add_run(s(subtitle))
        set_run(r2, "Cambria", 13, bold=True)

    if s(emphasis):
        p3 = doc.add_paragraph()
        tight_paragraph(p3, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=0, after_pt=14, line_spacing=1)
        r3 = p3.add_run(s(emphasis))
        set_run(r3, "Times New Roman", 14, bold=True)


def cover_add_meta_table(
    doc: Document,
    *,
    items: List[Tuple[str, str]],
    key_fill_hex: str = "D9E2F3",
    border_hex: str = "A6A6A6",
    key_w_in: float = 2.0,
    val_w_in: float = 4.8,
) -> None:
    items = items or []
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.style = "Table Grid"

    set_table_fixed_layout(tbl)
    set_table_borders(tbl, color_hex=border_hex)
    set_table_width_exact(tbl, Inches(key_w_in) + Inches(val_w_in))

    tbl.columns[0].width = Inches(key_w_in)
    tbl.columns[1].width = Inches(val_w_in)

    for k, v in items:
        rr = tbl.add_row()
        set_row_cant_split(rr, cant_split=True)

        c0, c1 = rr.cells[0], rr.cells[1]
        c0.width = Inches(key_w_in)
        c1.width = Inches(val_w_in)

        c0.text = ""
        c1.text = ""

        shade_cell(c0, key_fill_hex)
        set_cell_margins(c0, 80, 80, 120, 120)
        set_cell_margins(c1, 80, 80, 120, 120)

        p0 = c0.paragraphs[0]
        tight_paragraph(p0, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p0.add_run(s(k)), "Times New Roman", 11, bold=True)

        p1 = c1.paragraphs[0]
        tight_paragraph(p1, align=WD_ALIGN_PARAGRAPH.LEFT, before_pt=0, after_pt=0, line_spacing=1)
        set_run(p1.add_run(s(v)), "Times New Roman", 11, bold=False)

    doc.add_paragraph("")


def cover_add_footer_note(doc: Document, *, text: str) -> None:
    if not s(text):
        return
    doc.add_paragraph("")
    p = doc.add_paragraph()
    tight_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before_pt=10, after_pt=0, line_spacing=1)
    r = p.add_run(s(text))
    set_run(r, "Times New Roman", 9, bold=False, color=RGBColor(89, 89, 89))
